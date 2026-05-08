"""
构建带嵌入图片的论文 docx 文件。

策略:
1. Pandoc 编译 markdown -> docx 基础版（每次都重新编译，确保用最新 markdown）
2. 用 python-docx 在指定位置插入 12 张主图
3. 添加图说与编号
"""
import os
import sys
os.environ.setdefault("KMP_DUPLICATE_LIB_OK", "TRUE")
# 设置 stdout 为 UTF-8（Windows）
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

import shutil
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT = Path(r"E:\Claude project\SCI5\WUDAPT-FUSE")
PAPER_DIR = PROJECT / "docs" / "paper"
FIGURES = PROJECT / "figures"

# Pandoc 路径（Windows 默认安装位置）
PANDOC = Path(r"C:/Users/13950/AppData/Local/Pandoc/pandoc.exe")
if not PANDOC.exists():
    PANDOC = Path(shutil.which("pandoc") or "pandoc")

source_md = PAPER_DIR / "PAPER_FULL.md"
source_docx = PAPER_DIR / "WUDAPT-FUSE_paper_v1.docx"
output_docx = PAPER_DIR / "WUDAPT-FUSE_paper_with_figures.docx"

# === 0. 用 pandoc 从最新 markdown 重新编译 docx ===
print(f"[1/3] 用 pandoc 编译 {source_md.name} → {source_docx.name}")
pandoc_cmd = [
    str(PANDOC),
    str(source_md),
    '-o', str(source_docx),
    '--from=markdown',
    '--to=docx',
    '--standalone',
]
# 若有自定义 reference 模板（控制字体/字号/页边距），优先用之
ref_doc = PAPER_DIR / 'reference.docx'
if ref_doc.exists():
    pandoc_cmd += ['--reference-doc', str(ref_doc)]

result = subprocess.run(
    pandoc_cmd,
    capture_output=True, text=True, encoding='utf-8', errors='replace',
)
if result.returncode != 0:
    print(f"  ✗ pandoc 失败 (exit {result.returncode}):")
    print(result.stderr or result.stdout)
    sys.exit(1)
print(f"  ✓ 基础 docx: {source_docx.stat().st_size/1024:.1f} KB")

# 1. 复制基础 docx 作为目标，再插入图片
shutil.copy(source_docx, output_docx)
doc = Document(output_docx)

# 2. 找到 # Figures 段落，把对应的图插入到每个 ## Figure X 标题后
fig_files = {
    'Figure 1': FIGURES / "fig01_concept.png",                # 研究区 + 输入数据 (3 maps)
    'Figure 2': FIGURES / "fig02_framework.png",              # WUDAPT-FUSE 框架图（独立）
    'Figure 3': FIGURES / "fig03_utci_inequality.png",        # UTCI dataset & baseline inequality
    'Figure 4': FIGURES / "fig04_lcz_subcategories.png",      # LCZ subcategories from morphology
    'Figure 5': FIGURES / "fig05_subcategory_gini_test.png",  # Subcategorization fails (key)
    'Figure 6': FIGURES / "fig06_tda_analysis.png",           # TDA
    'Figure 7': FIGURES / "fig07_hybrid_clustering.png",      # Hybrid form + topology
    'Figure 8': FIGURES / "fig08_nonspatial_test.png",        # Geographic position breaks 5 pp
    'Figure 9': FIGURES / "fig09_geography_mechanism.png",    # Coast + mountain mechanism
    'Figure 10': FIGURES / "fig10_diurnal_mechanism.png",     # Diurnal validation
    'Figure 11': FIGURES / "fig11_seasonal_robustness.png",   # Seasonal robustness
    'Figure 12': FIGURES / "fig12_multicity_validation.png",  # Cross-city
}

# 3. 收集每个 Figure 标题段落 OBJECT (object 引用稳定，索引会因插入而失效)
fig_paragraphs = {}
for p in doc.paragraphs:
    text = p.text.strip()
    for fig_name in fig_files.keys():
        if text.startswith(f"{fig_name}.") and p.style.name == "Heading 2":
            fig_paragraphs[fig_name] = p  # 存对象，不存索引

print(f"找到 {len(fig_paragraphs)} 个 Figure 标题段落: "
      f"{sorted(fig_paragraphs.keys(), key=lambda x: int(x.split()[-1]))}")

# 4. 在每个 Figure 标题后插入图（倒序以保持索引有效）
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def insert_image_after(paragraph, image_path, width_inches=6.0):
    """在指定段落后插入一张图（用 docx 标准 API）"""
    # 在 paragraph 之后创建新段落
    new_p_xml = OxmlElement("w:p")
    paragraph._element.addnext(new_p_xml)

    # 用 docx Paragraph 包装并添加图片
    from docx.text.paragraph import Paragraph as DocxParagraph
    new_para = DocxParagraph(new_p_xml, paragraph._parent)

    run = new_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    # 居中（通过 pPr -> jc）
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    new_p_xml.insert(0, pPr)

    return new_para

# 顺序无关：因为 fig_paragraphs[name] 存的是对象引用，
# 插入只通过 paragraph._element.addnext() 操作 XML 节点，
# 不依赖 doc.paragraphs 列表索引，所以正向/逆向都安全。
fig_order = [f'Figure {n}' for n in range(1, 13)]  # Figure 1 到 Figure 12
for fig_name in fig_order:
    if fig_name in fig_paragraphs:
        para = fig_paragraphs[fig_name]  # 直接用对象引用，绝不会过期
        img_path = fig_files[fig_name]
        if img_path.exists():
            insert_image_after(para, img_path, width_inches=6.5)
            print(f"  ✓ 插入 {fig_name}: {img_path.name}")
        else:
            print(f"  ✗ 找不到 {img_path}")
    else:
        print(f"  ⚠ docx 中未找到 '{fig_name}.' 标题段落")

# 5. 保存
doc.save(output_docx)
size_kb = output_docx.stat().st_size / 1024
print(f"\n📄 输出: {output_docx.name}")
print(f"   大小: {size_kb:.1f} KB ({size_kb/1024:.2f} MB)")
