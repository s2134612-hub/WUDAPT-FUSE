"""验证 docx 论文文件结构"""
import os
os.environ.setdefault("KMP_DUPLICATE_LIB_OK", "TRUE")

from docx import Document
from pathlib import Path

doc_path = Path(r"E:\Claude project\SCI5\WUDAPT-FUSE\docs\paper\WUDAPT-FUSE_paper_v1.docx")
doc = Document(doc_path)

print(f"📄 Paper docx: {doc_path.name}")
print(f"   File size: {doc_path.stat().st_size / 1024:.1f} KB")
print(f"   Paragraphs: {len(doc.paragraphs)}")
print(f"   Tables:     {len(doc.tables)}")

# 统计章节
print("\n📑 Section structure:")
section_counts = {}
for p in doc.paragraphs:
    style = p.style.name
    section_counts[style] = section_counts.get(style, 0) + 1
for style, count in sorted(section_counts.items(), key=lambda x: -x[1]):
    print(f"   {style:30} {count}")

# 字数统计
total_words = 0
for p in doc.paragraphs:
    total_words += len(p.text.split())
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            total_words += len(cell.text.split())
print(f"\n📊 Word count: {total_words:,}")

# 显示标题层级
print("\n📚 Headings (Top-level):")
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        level = int(p.style.name.replace('Heading ', '')) if 'Heading ' in p.style.name else 0
        indent = "  " * (level - 1)
        text = p.text[:60]
        print(f"   {indent}H{level}: {text}")
