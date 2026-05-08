"""
将 SCS 投稿 docx 应用 Sustainable Cities and Society / Elsevier 投稿格式：

  ✓ 字体: Times New Roman 12 pt (正文); 表格 11 pt; 标题加粗
  ✓ 行距: 双倍 (2.0) — 评审版必需
  ✓ 段后距: 0 pt (双倍行距下不需要额外段后距)
  ✓ 段前距: 0 pt
  ✓ 首行缩进: 0 (英文期刊 block style 标准)
  ✓ 对齐: 左对齐 (ragged right)
  ✓ 页边距: 2.5 cm 上下左右
  ✓ 连续行号: 全文从 1 开始 (Elsevier 评审强制)
  ✓ 页码: 居中页脚

Usage:
  python scripts/apply_scs_formatting.py
"""
import os
os.environ.setdefault("KMP_DUPLICATE_LIB_OK", "TRUE")

import sys
import shutil
from pathlib import Path

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


PROJECT = Path(r"E:\Claude project\SCI5\WUDAPT-FUSE")
PAPER_DIR = PROJECT / "docs" / "paper"
SRC_DOCX = PAPER_DIR / "WUDAPT-FUSE_paper_with_figures.docx"
OUT_DOCX = PAPER_DIR / "WUDAPT-FUSE_SCS_submission.docx"

# === SCS 格式参数 ===
BODY_FONT = "Times New Roman"
BODY_SIZE_PT = 12
TABLE_SIZE_PT = 11
HEADING_BASE_SIZE_PT = 12   # 与正文同字号，加粗
LINE_SPACING_DOUBLE = 2.0   # SCS 必需
MARGIN_CM = 2.5             # 2.5 cm 全边距
LINE_NUMBER_RESTART = "continuous"  # 全文连续


def set_font_run(run, name=BODY_FONT, size_pt=BODY_SIZE_PT,
                  bold=None, italic=None):
    """设置 run 字体（覆盖 style 的设置）"""
    run.font.name = name
    # 设置中文字体（防止中文回退）
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:eastAsia'), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def set_paragraph_format(p, line_spacing=LINE_SPACING_DOUBLE,
                          space_before=0, space_after=0,
                          first_line_indent=0,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """设置段落格式（行距/段距/缩进/对齐）"""
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.first_line_indent = Pt(first_line_indent)
    if alignment is not None:
        pf.alignment = alignment


def apply_body_format(doc):
    """对所有正文段落应用 Times New Roman 12pt + 双倍行距"""
    body_styles = ('Normal', 'First Paragraph', 'Body Text', 'Compact')
    heading_styles = (f'Heading {i}' for i in range(1, 7))

    n_body = 0
    n_heading = 0
    n_caption = 0

    for p in doc.paragraphs:
        style = p.style.name

        if style.startswith('Heading '):
            # 标题: 12 pt 粗体, 但要有上下间距
            for run in p.runs:
                set_font_run(run, BODY_FONT, BODY_SIZE_PT, bold=True)
            # 标题级别决定空白间距
            level = int(style.split()[-1]) if style.split()[-1].isdigit() else 1
            sb = max(12, 6 * (5 - level))   # H1 24pt, H2 18pt, H3 12pt
            sa = max(6, 3 * (5 - level))    # H1 12pt, H2 9pt, H3 6pt
            set_paragraph_format(p,
                                  line_spacing=LINE_SPACING_DOUBLE,
                                  space_before=sb, space_after=sa,
                                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
            n_heading += 1

        elif style == 'Caption':
            # Figure/Table caption: 11 pt italic
            for run in p.runs:
                set_font_run(run, BODY_FONT, 11, italic=True)
            set_paragraph_format(p,
                                  line_spacing=1.15,
                                  space_before=3, space_after=6,
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
            n_caption += 1

        elif style in body_styles:
            for run in p.runs:
                set_font_run(run, BODY_FONT, BODY_SIZE_PT)
            set_paragraph_format(p,
                                  line_spacing=LINE_SPACING_DOUBLE,
                                  space_before=0, space_after=0,
                                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
            n_body += 1

        else:
            # 兜底: 列表项、引用等也用同样字体
            for run in p.runs:
                set_font_run(run, BODY_FONT, BODY_SIZE_PT)
            # 列表项保留单倍行距以避免过于稀疏
            if 'List' in style:
                set_paragraph_format(p,
                                      line_spacing=LINE_SPACING_DOUBLE,
                                      space_before=0, space_after=0,
                                      alignment=WD_ALIGN_PARAGRAPH.LEFT)
            n_body += 1

    return n_body, n_heading, n_caption


def apply_table_format(doc):
    """表格内字体: Times New Roman 11 pt, 单倍行距"""
    n_cells = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_font_run(run, BODY_FONT, TABLE_SIZE_PT)
                    pf = p.paragraph_format
                    pf.line_spacing = 1.15
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    n_cells += 1
    return n_cells


def set_margins(doc, margin_cm=MARGIN_CM):
    """设置全文页边距"""
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)


def add_line_numbers(doc, restart='continuous', distance_pt=18):
    """添加连续行号 (SCS 评审强制要求)"""
    for section in doc.sections:
        sectPr = section._sectPr
        # 删除旧的 lnNumType（避免重复）
        for old in sectPr.findall(qn('w:lnNumType')):
            sectPr.remove(old)
        lnNumType = OxmlElement('w:lnNumType')
        lnNumType.set(qn('w:countBy'), '1')         # 每行一个号
        lnNumType.set(qn('w:start'), '1')           # 从 1 开始
        lnNumType.set(qn('w:distance'), str(distance_pt * 20))  # 距离左边距 (twips)
        lnNumType.set(qn('w:restart'), restart)     # continuous = 全文连续
        sectPr.append(lnNumType)


def add_page_numbers(doc):
    """页脚居中页码"""
    for section in doc.sections:
        footer = section.footer
        # 清空已有内容
        for p in footer.paragraphs:
            p._element.getparent().remove(p._element)

        # 添加新段落
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 插入 PAGE 字段
        run = p.add_run()
        set_font_run(run, BODY_FONT, BODY_SIZE_PT)

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)


def update_default_styles(doc):
    """更新 Normal/Body Text/Heading 样式的字体 + 段落格式"""
    body_styles = ['Normal', 'Body Text', 'First Paragraph', 'Compact']
    heading_styles = ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']

    for style_name in body_styles + heading_styles:
        try:
            style = doc.styles[style_name]
            # 字体
            style.font.name = BODY_FONT
            style.font.size = Pt(BODY_SIZE_PT)
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), BODY_FONT)
            rFonts.set(qn('w:hAnsi'), BODY_FONT)
            rFonts.set(qn('w:cs'), BODY_FONT)
            rFonts.set(qn('w:eastAsia'), BODY_FONT)

            # 段落格式: 重置默认间距为 0, 双倍行距
            spf = style.paragraph_format
            if style_name in body_styles:
                spf.line_spacing = LINE_SPACING_DOUBLE
                spf.space_before = Pt(0)
                spf.space_after = Pt(0)
                spf.first_line_indent = Pt(0)
            elif style_name in heading_styles:
                level = int(style_name.split()[-1])
                spf.line_spacing = LINE_SPACING_DOUBLE
                spf.space_before = Pt(max(12, 6 * (5 - level)))
                spf.space_after = Pt(max(6, 3 * (5 - level)))
                # 标题加粗
                style.font.bold = True
        except KeyError:
            pass


def main():
    if not SRC_DOCX.exists():
        print(f"❌ 源文件不存在: {SRC_DOCX}")
        return 1

    print(f"读取: {SRC_DOCX.name}")
    shutil.copy(SRC_DOCX, OUT_DOCX)
    doc = Document(OUT_DOCX)

    print("\n[1/5] 更新默认样式 (Normal/Heading) → Times New Roman 12pt")
    update_default_styles(doc)

    print("[2/5] 设置正文字体 + 双倍行距")
    n_body, n_heading, n_caption = apply_body_format(doc)
    print(f"      正文段落: {n_body} | 标题: {n_heading} | 图表说明: {n_caption}")

    print("[3/5] 设置表格字体 (TNR 11pt)")
    n_cells = apply_table_format(doc)
    print(f"      表格 cell: {n_cells}")

    print(f"[4/5] 设置页边距 ({MARGIN_CM} cm 全边距)")
    set_margins(doc, margin_cm=MARGIN_CM)

    print("[5/5] 添加连续行号 + 居中页码")
    add_line_numbers(doc, restart='continuous')
    add_page_numbers(doc)

    print(f"\n保存: {OUT_DOCX.name}")
    doc.save(OUT_DOCX)
    size_kb = OUT_DOCX.stat().st_size / 1024
    print(f"   大小: {size_kb:.1f} KB ({size_kb/1024:.2f} MB)")

    return 0


if __name__ == "__main__":
    sys.exit(main())
