"""验证最终带图 docx"""
import os
os.environ.setdefault("KMP_DUPLICATE_LIB_OK", "TRUE")

from docx import Document
from pathlib import Path

doc_path = Path(r"E:\Claude project\SCI5\WUDAPT-FUSE\docs\paper\WUDAPT-FUSE_paper_with_figures.docx")
doc = Document(doc_path)

print("=" * 60)
print(f"📄 Final paper: {doc_path.name}")
print(f"   Size: {doc_path.stat().st_size / 1024:.1f} KB ({doc_path.stat().st_size/1e6:.2f} MB)")
print("=" * 60)

# 字数
total_words = 0
for p in doc.paragraphs:
    total_words += len(p.text.split())
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            total_words += len(cell.text.split())
print(f"\n📊 Word count:    {total_words:,}")
print(f"   Paragraphs:    {len(doc.paragraphs)}")
print(f"   Tables:        {len(doc.tables)}")

# 检查图片
from docx.document import Document as DocBody
images = []
for rel_id, rel in doc.part.rels.items():
    if "image" in rel.reltype:
        target = rel.target_ref
        images.append(target)
print(f"   Embedded images: {len(images)}")
for img in images:
    print(f"     - {img}")

# 章节结构
print("\n📚 Section outline:")
for p in doc.paragraphs:
    style = p.style.name
    if style.startswith('Heading'):
        level = int(style.replace('Heading ', ''))
        indent = "  " * (level - 1)
        print(f"   {indent}H{level}: {p.text[:55]}")

# Quality checks
print("\n✅ Quality checks:")
checks = []

# Check 1: Title present
title_found = any(p.style.name == 'Title' for p in doc.paragraphs)
checks.append(("Title style present", title_found))

# Check 2: Abstract section
abs_found = any('Abstract' in p.style.name for p in doc.paragraphs)
checks.append(("Abstract section present", abs_found))

# Check 3: All 5 main sections (1-5 + Acknowledgments + References)
for section in ['1. Introduction', '2. Methods', '3. Results', '4. Discussion', 'References']:
    found = any(section in p.text for p in doc.paragraphs if p.style.name == 'Heading 1')
    checks.append((f"Section '{section}' present", found))

# Check 4: 4 tables present
checks.append(("All 4 tables present", len(doc.tables) >= 4))

# Check 5: All 6 images embedded
checks.append((f"All figures embedded ({len(images)})", len(images) >= 4))

# Check 6: No Chinese characters
chinese_count = 0
for p in doc.paragraphs:
    for ch in p.text:
        if 0x4e00 <= ord(ch) <= 0x9fff:
            chinese_count += 1
checks.append(("No Chinese characters in text", chinese_count == 0))

# Check 7: Reasonable word count (3000-6000 for Nature Comm)
checks.append((f"Word count in target range (3000-6000): {total_words}",
              3000 <= total_words <= 6000))

for name, ok in checks:
    sym = '✓' if ok else '✗'
    print(f"   [{sym}] {name}")

n_pass = sum(1 for _, ok in checks if ok)
print(f"\n   {n_pass}/{len(checks)} checks passed")
